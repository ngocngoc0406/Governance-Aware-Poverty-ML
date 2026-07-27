import os
import glob
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_markdown_to_doc(doc, md_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.readlines()
        
    for line in content:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('![') and ']' in line and '(' in line:
            # Handle image embedding
            try:
                # Extract image path: ![alt](path)
                img_path = line.split('(')[1].split(')')[0].replace('file:///', '').replace('C:/', 'C:\\').replace('/', '\\')
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(6.0))
                else:
                    doc.add_paragraph(f"[Image missing: {img_path}]")
            except Exception as e:
                doc.add_paragraph(f"[Error loading image: {line}]")
        elif '|' in line and '-' in line and line.count('|') > 2:
            # Skip markdown table separators like |---|---|
            continue
        elif line.startswith('|') and line.endswith('|'):
            # Very basic table handler - just convert to text separated by tabs for now
            cells = [c.strip() for c in line.split('|')[1:-1]]
            doc.add_paragraph(' \t '.join(cells))
        else:
            # Normal paragraph
            # Remove basic markdown bold/italics markers for a cleaner look
            clean_line = line.replace('**', '').replace('*', '')
            doc.add_paragraph(clean_line)
            
    doc.add_page_break()

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Explainable Machine Learning for Transparent Poverty Targeting: A Comparative Study', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add files in order
    md_files = [
        'paper_draft/01_Introduction.md',
        'paper_draft/02_Literature_Review.md',
        'paper_draft/03_Methodology.md',
        'paper_draft/04_Results.md',
        'paper_draft/05_Discussion_Conclusion.md'
    ]
    
    for md_file in md_files:
        if os.path.exists(md_file):
            add_markdown_to_doc(doc, md_file)
            
    doc.save('Paper_Final_Draft.docx')
    print("Paper_Final_Draft.docx has been generated successfully.")

if __name__ == '__main__':
    main()
